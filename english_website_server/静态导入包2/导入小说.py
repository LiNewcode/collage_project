import re

from docx import Document
import docx
document = Document("F:\PyCourseDemo\end_of_term_work\静态导入包2\小说包\爱丽丝漫游奇境记 .docx")

# 获取文档所有表格
all_tables = document.tables
# print(document.paragraphs[0])
content = []
for p in document.paragraphs:
    # print(p.text)
    content.append(p.text)
print(content)
print(len(re.findall('\n\s(\d{1,2})\s',''.join(content),re.S)))
paragraph = re.split('\n\s(\d{1,2})\s',''.join(content),re.S)
print(len(paragraph))
newList = []
for item in paragraph:
    if len(item)<4:
        print(item)
    else:
        newList.append(item)
print(newList[1::2])
print(newList[::2])

#打印all_tables类型
# print(all_tables) #得到<class 'list'>，即列表
#开始循环读取表格列表
# for table in all_tables:
#     print(table)
# #循环读取表格的每一行
# for row in table.rows:
#     #print(row)
#     #循环读取表格的每一个单元格
#     for cell in row.cells:
#         #打印单元格里的内容
#         print(cell.text) #打印
