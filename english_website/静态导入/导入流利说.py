import re
from glob import glob

import docx

from end_of_term_work.Tool.link_db import db_obj


class import_lls():
    def __init__(self, path):
        self.path = path

    def run(self, _id):
        db = db_obj(col='e_pee')
        kaoyan = db.search({}, {'_id': 0})
        db.close()
        pee = []
        for i in kaoyan:
            pee.append(i['word'][0])
        # print(pee)
        file = docx.Document(self.path)
        # print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

        en_t = []
        en_w = []
        ch_t = []
        ch_w = []
        end = 0
        state = False
        # 输出每一段的内容
        text = []
        for i in range(len(file.paragraphs)):
            # print("第"+str(i)+"段的内容是："+text[i])
            if file.paragraphs[i].text != '':
                text.append(file.paragraphs[i].text)
        for i in range(len(text)):
            txt = text[i]
            if i + 1 < len(text):
                txt_next = text[i + 1]
            if len(re.findall('正文', txt)) > 0:
                state = True
            if len(re.findall('重点词汇', txt)) > 0:
                end = i
                state = False
            if state and len(re.findall('[\u4e00-\u9fa5]', txt)) > 0 and len(
                    re.findall('[\u4e00-\u9fa5]', txt_next)) <= 0:
                en_t.append(i + 1)
                ch_w.append(i)
            if state and len(re.findall('[\u4e00-\u9fa5]', txt)) <= 0 and len(
                    re.findall('[\u4e00-\u9fa5]', txt_next)) > 0:
                en_w.append(i)
                ch_t.append(i + 1)
        if end == 0:
            print('end出错')
            print(self.path)
            for i in range(len(text)):
                txt = text[i]
                if i + 1 < len(text):
                    txt_next = text[i + 1]
                if len(re.findall('英文原文', txt)) > 0:
                    state = True
                if len(re.findall('内容拓展', txt)) > 0:
                    end = i
                    state = False
                if state and len(re.findall('[\u4e00-\u9fa5]', txt)) > 0 and len(
                        re.findall('[\u4e00-\u9fa5]', txt_next)) <= 0:
                    en_t.append(i + 1)
                    ch_w.append(i)
                if state and len(re.findall('[\u4e00-\u9fa5]', txt)) <= 0 and len(
                        re.findall('[\u4e00-\u9fa5]', txt_next)) > 0:
                    en_w.append(i)
                    ch_t.append(i + 1)
            return
        ch_w.append(end)
        annotation = text[end:]
        # print(annotation)
        # else:
        #     print("[ 英文 ]第" + str(i) + "段的内容是：" + text[i])
        # print('[en_t]:', en_t)
        # print('[en_w]:', en_w)
        # print('[ch_t]:', ch_t)
        # print('[ch_w]:', ch_w)
        index = 0

        if en_t[0] == en_w[0]:
            title = text[en_t[0]]
        elif en_t[0] + 1 == en_w[0]:
            title = text[en_t[0]] + " " + text[en_w[0]]
        else:
            print(en_t, text[en_t[0]])
            print(en_w, text[en_w[0]])
            print(self.path)
            return
            # print('title:', title)
        index = 1
        article = []
        article_pee = []
        while index < len(en_t):
            paragraph = index
            content = ''.join(text[en_t[index]:en_w[index] + 1])
            des = ''.join(text[ch_t[index]:ch_w[index + 1] + 1])
            content = re.split('("|\'|\s|,|\.|\d{1,2})', content)
            # print(content)
            pc = []
            for i in content:
                isPee = False
                if i.isalpha():
                    if i in pee:
                        isPee = True
                        article_pee.append(i)
                        item = {'word': i, 'signal': 'en', 'isPee': isPee}
                    else:
                        item = {'word': i, 'signal': 'en', 'isPee': isPee}
                else:
                    item = {'word': i, 'signal': 'other', 'isPee': isPee}
                pc.append(item)
            paragraph = {'paragraph': index, 'content': pc, 'des': des}
            # print(paragraph)
            article.append(paragraph)
            index += 1
        if len(re.findall('[\u4e00-\u9fa5]', text[ch_t[0] + 1])) > 0:
            title_des = text[ch_t[0]] + text[ch_t[0] + 1]
        else:
            title_des = text[ch_t[0]]

        content = {'_id': _id, 'title': title, 'title_des': title_des, 'article': article, 'annotation': annotation,
                   'article_pee': article_pee}
        # print(content)
        db = db_obj(col='lls')
        db.update(find={'_id': _id}, content=content)
        db.close()
